from flask import Flask, request, render_template, send_file
from transformers import pipeline
from utils.file_extractor import extract_text_from_file
from utils.nlp_processor import process_text
from utils.question_generator import generate_questions_with_answers
import asyncio
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
import random
import os

app = Flask(__name__)

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        file = request.files['file']
        quiz_type = request.form.get('quiz_type')
        num_questions = int(request.form.get('num_questions'))
        quiz_level = request.form.get('quiz_level')

        if file and (file.filename.endswith(('.docx', '.pdf', '.txt'))):
            content = extract_text_from_file(file)
            # Store the original file name (without extension) to use later
            original_filename = os.path.splitext(file.filename)[0]
            questions = asyncio.run(generate_questions_with_answers(content, quiz_type, num_questions))  # Pass num_questions
            return render_template('quiz_preview.html', questions=questions, original_filename=original_filename)

    return render_template('index.html')

@app.route('/download_quiz', methods=['POST'])
def download_quiz():
    questions = request.json.get('quiz_data', [])
    file_format = request.json.get('file_format', 'pdf')
    include_answers = request.json.get('include_answers', True)
    original_filename = request.json.get('original_filename', 'quiz')  # Get original filename from request

    print(f"Download requested: {file_format}, Filename: {original_filename}")

    if file_format == 'pdf':
        return generate_pdf(questions, include_answers, original_filename)
    elif file_format == 'docx':
        return generate_docx(questions, include_answers, original_filename)

    return "Invalid format", 400

def generate_pdf(questions, include_answers, original_filename):
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    y = height - 50

    pdf.setFont("Helvetica", 12)

    for index, question_data in enumerate(questions, start=1):
        question = question_data['question']
        choices = question_data.get('choices', [])
        answer = question_data['answer']

        pdf.drawString(50, y, f"Q{index}: {question}")
        y -= 20
        
        # Check if choices is not None and has elements
        if choices:
            for i, choice in enumerate(choices, start=1):
                pdf.drawString(70, y, f"{i}. {choice}")
                y -= 20

        if include_answers:
            pdf.drawString(50, y, f"Answer: {answer}")
            y -= 20

        y -= 20  # Add some space before the next question

        if y <= 50:  # Create a new page if the content reaches the bottom
            pdf.showPage()
            y = height - 50

    pdf.save()
    buffer.seek(0)

    # Create a new file name using the original file name
    download_name = f"{original_filename}_quiz.pdf"
    return send_file(buffer, as_attachment=True, download_name=download_name, mimetype='application/pdf')

def generate_docx(questions, include_answers, original_filename):
    doc = Document()

    for index, question_data in enumerate(questions, start=1):
        question = question_data['question']
        choices = question_data.get('choices', [])
        answer = question_data['answer']

        doc.add_paragraph(f"Q{index}: {question}")

        # Check if choices is not None and has elements
        if choices:
            for i, choice in enumerate(choices, start=1):
                doc.add_paragraph(f"{i}. {choice}", style='ListNumber')

        if include_answers:
            doc.add_paragraph(f"Answer: {answer}")

        doc.add_paragraph("")  # Add some space

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Create a new file name using the original file name
    download_name = f"{original_filename}_quiz.docx"
    return send_file(buffer, as_attachment=True, download_name=download_name, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

if __name__ == '__main__':
    app.run(debug=True)
